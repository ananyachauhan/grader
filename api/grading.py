"""
Grading API endpoints
"""
from flask import Blueprint, jsonify, request, send_file, abort
import sys
import os
import json
import importlib.util
from pathlib import Path
from datetime import datetime
from dotenv import load_dotenv
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False

# Get project root and scripts directory
# __file__ is at api/grading.py, so:
# .parent = api/
# .parent.parent = grader/ (project root)
project_root = Path(__file__).parent.parent

# Load environment variables from project root .env file
env_path = project_root / '.env'
load_dotenv(dotenv_path=env_path)
scripts_dir = project_root / 'scripts'

# Import grading_workflow module directly
grading_workflow_path = scripts_dir / 'grading_workflow.py'
spec1 = importlib.util.spec_from_file_location("grading_workflow", grading_workflow_path)
grading_workflow = importlib.util.module_from_spec(spec1)
spec1.loader.exec_module(grading_workflow)
grade_document = grading_workflow.grade_document
load_config = grading_workflow.load_config

# Import ai_grader module directly
ai_grader_path = scripts_dir / 'ai_grader.py'
spec2 = importlib.util.spec_from_file_location("ai_grader", ai_grader_path)
ai_grader = importlib.util.module_from_spec(spec2)
spec2.loader.exec_module(ai_grader)
load_rubric = ai_grader.load_rubric

grading_bp = Blueprint('grading', __name__)

@grading_bp.route('/rubrics', methods=['GET'])
def list_rubrics():
    """List all available rubrics"""
    rubrics_dir = Path(__file__).parent.parent / 'rubrics'
    rubrics = []
    
    if rubrics_dir.exists():
        for rubric_file in rubrics_dir.glob('*.json'):
            try:
                rubric = load_rubric(str(rubric_file))
                rubrics.append({
                    'name': rubric.get('name', rubric_file.stem),
                    'filename': rubric_file.name,
                    'total_points': rubric.get('total_points', 0),
                    'criteria_count': len(rubric.get('criteria', []))
                })
            except:
                pass
    
    return jsonify({'rubrics': rubrics})

@grading_bp.route('/rubrics/<filename>', methods=['GET'])
def get_rubric(filename):
    """Get full rubric data by filename"""
    # Security: ensure filename doesn't contain path traversal
    if '..' in filename or '/' in filename or '\\' in filename:
        return jsonify({'error': 'Invalid filename'}), 400
    
    rubrics_dir = Path(__file__).parent.parent / 'rubrics'
    rubric_path = rubrics_dir / filename
    
    if not rubric_path.exists():
        return jsonify({'error': 'Rubric not found'}), 404
    
    try:
        rubric = load_rubric(str(rubric_path))
        return jsonify({'rubric': rubric})
    except Exception as e:
        return jsonify({'error': f'Error loading rubric: {str(e)}'}), 500

@grading_bp.route('/rubrics/<filename>/original', methods=['GET'])
def get_rubric_original(filename):
    """Serve the original rubric file (Word/PDF) for embedding"""
    import logging
    logger = logging.getLogger(__name__)
    
    # Security: ensure filename doesn't contain path traversal
    if '..' in filename or '/' in filename or '\\' in filename:
        abort(400)
    
    rubrics_dir = Path(__file__).parent.parent / 'rubrics'
    
    # First, try to find the original file by looking up the JSON
    json_filename = filename
    if not json_filename.endswith('.json'):
        json_filename = f"{filename}.json"
    
    json_path = rubrics_dir / json_filename
    
    if json_path.exists():
        # Load JSON to get original file path
        try:
            rubric = load_rubric(str(json_path))
            original_file_path = rubric.get('_original_file_path')
            logger.info(f"Looking for original file: {original_file_path}")
            if original_file_path:
                original_path = rubrics_dir / original_file_path
                logger.info(f"Checking if file exists: {original_path}, exists: {original_path.exists()}")
                if original_path.exists():
                    # Determine MIME type
                    ext = original_path.suffix.lower()
                    mimetypes = {
                        '.pdf': 'application/pdf',
                        '.doc': 'application/msword',
                        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    }
                    mimetype = mimetypes.get(ext, 'application/octet-stream')
                    logger.info(f"Serving file: {original_path} with mimetype: {mimetype}")
                    return send_file(str(original_path), mimetype=mimetype)
        except Exception as e:
            logger.error(f"Error loading rubric JSON: {e}")
            import traceback
            traceback.print_exc()
    
    # Fallback: try to find file with same base name but different extension
    base_name = Path(filename).stem
    logger.info(f"Trying fallback with base name: {base_name}")
    for ext in ['.pdf', '.doc', '.docx']:
        original_path = rubrics_dir / f"{base_name}{ext}"
        logger.info(f"Checking fallback path: {original_path}, exists: {original_path.exists()}")
        if original_path.exists():
            mimetypes = {
                '.pdf': 'application/pdf',
                '.doc': 'application/msword',
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            }
            mimetype = mimetypes.get(ext, 'application/octet-stream')
            logger.info(f"Serving fallback file: {original_path} with mimetype: {mimetype}")
            return send_file(str(original_path), mimetype=mimetype)
    
    logger.error(f"Original file not found for rubric: {filename}")
    abort(404)

@grading_bp.route('/rubrics/<filename>', methods=['DELETE'])
def delete_rubric(filename):
    """Delete a rubric file"""
    # Security: ensure filename doesn't contain path traversal
    if '..' in filename or '/' in filename or '\\' in filename:
        return jsonify({'error': 'Invalid filename'}), 400
    
    rubrics_dir = Path(__file__).parent.parent / 'rubrics'
    rubric_path = rubrics_dir / filename
    
    # Only allow deletion of JSON files in rubrics directory
    if not filename.endswith('.json'):
        return jsonify({'error': 'Can only delete JSON rubric files'}), 400
    
    if not rubric_path.exists():
        return jsonify({'error': 'Rubric not found'}), 404
    
    # Prevent deletion of default rubric if it exists
    # (optional safety check)
    
    try:
        # Load rubric to get name for response
        try:
            rubric = load_rubric(str(rubric_path))
            rubric_name = rubric.get('name', filename)
        except:
            rubric_name = filename
        
        # Delete the file
        rubric_path.unlink()
        
        return jsonify({
            'success': True,
            'message': f'Rubric "{rubric_name}" deleted successfully'
        })
    except Exception as e:
        return jsonify({'error': f'Error deleting rubric: {str(e)}'}), 500

@grading_bp.route('/grade', methods=['POST'])
def grade_single():
    """Grade a single document"""
    data = request.json
    doc_id = data.get('doc_id')
    rubric_filename = data.get('rubric_filename')
    custom_instructions = data.get('custom_instructions', '').strip() or None
    
    if not doc_id:
        return jsonify({'error': 'doc_id is required'}), 400
    
    if not rubric_filename:
        return jsonify({'error': 'rubric_filename is required'}), 400
    
    # Get rubric path
    rubric_path = Path(__file__).parent.parent / 'rubrics' / rubric_filename
    if not rubric_path.exists():
        return jsonify({'error': f'Rubric not found: {rubric_filename}'}), 404
    
    try:
        config = load_config()
        result = grade_document(doc_id, str(rubric_path), config, custom_instructions)
        return jsonify(result)
    except Exception as e:
        return jsonify({'error': str(e), 'success': False}), 500

@grading_bp.route('/grade/batch', methods=['POST'])
def grade_batch():
    """Grade multiple documents"""
    data = request.json
    doc_ids = data.get('doc_ids', [])
    doc_types = data.get('doc_types', {})  # {doc_id: True if Word doc, False if Google Doc}
    rubric_filename = data.get('rubric_filename')
    custom_instructions = data.get('custom_instructions', '').strip() or None
    
    if not doc_ids:
        return jsonify({'error': 'doc_ids is required'}), 400
    
    if not rubric_filename:
        return jsonify({'error': 'rubric_filename is required'}), 400
    
    rubric_path = Path(__file__).parent.parent / 'rubrics' / rubric_filename
    if not rubric_path.exists():
        return jsonify({'error': f'Rubric not found: {rubric_filename}'}), 404
    
    config = load_config()
    results = []
    
    for doc_id in doc_ids:
        is_word_doc = doc_types.get(doc_id, False)
        try:
            result = grade_document(doc_id, str(rubric_path), config, custom_instructions, is_word_doc=is_word_doc)
            results.append({
                'doc_id': doc_id,
                'success': result.get('success', False),
                'total_score': result.get('total_score', 0),
                'scores': result.get('scores', {}),
                'strengths': result.get('strengths', ''),
                'key_issues': result.get('key_issues', ''),
                'suggestions': result.get('suggestions', ''),
                'error': result.get('error'),
                'converted_doc_id': result.get('converted_doc_id'),  # New Google Doc ID if converted
                'original_doc_id': result.get('original_doc_id'),  # Original Word doc ID if converted
                'message': result.get('message', '')
            })
        except Exception as e:
            results.append({
                'doc_id': doc_id,
                'success': False,
                'error': str(e)
            })
    
    return jsonify({'results': results})


def parse_rubric_with_ai(document_text: str) -> dict:
    """
    Use AI to intelligently parse rubric from document text.
    This handles any format - tables, lists, paragraphs, etc.
    """
    if not GEMINI_AVAILABLE:
        raise ValueError("google-generativeai library not installed. Install with: pip install google-generativeai")
    
    api_key = os.getenv('GEMINI_API_KEY')
    if not api_key:
        # Try to reload .env file in case it wasn't loaded
        env_path = Path(__file__).parent.parent / '.env'
        if env_path.exists():
            load_dotenv(dotenv_path=env_path, override=True)
            api_key = os.getenv('GEMINI_API_KEY')
    
    if not api_key:
        raise ValueError("GEMINI_API_KEY environment variable not set. Please add GEMINI_API_KEY=your_key to your .env file in the project root.")
    
    # Strip any whitespace that might have been accidentally included
    api_key = api_key.strip()
    
    genai.configure(api_key=api_key)
    
    # Configure model
    generation_config = {
        "temperature": 0.1,  # Low temperature for consistent structured output
        "top_p": 0.95,
        "top_k": 40,
        "max_output_tokens": 4000,  # Increased to handle longer rubrics
    }
    
    # First, try to find an available model that supports generateContent
    model = None
    model_name_to_use = None
    
    try:
        # List all available models
        available_models = []
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                available_models.append(m.name)
        
        # Try to find a working model
        # Priority order: gemini-2.5-flash (fastest), gemini-pro-latest, gemini-2.5-pro, or any available
        preferred_models = [
            "models/gemini-2.5-flash",  # Latest fast model
            "models/gemini-flash-latest",  # Latest flash
            "models/gemini-pro-latest",  # Latest pro
            "models/gemini-2.5-pro",  # Latest pro version
            "models/gemini-2.0-flash",  # Alternative flash
        ]
        
        for preferred in preferred_models:
            # Check if this model is in the available list
            if preferred in available_models:
                model_name_to_use = preferred
                break
        
        # If no preferred model found, use the first available one
        if not model_name_to_use and available_models:
            model_name_to_use = available_models[0]
        
        if not model_name_to_use:
            raise ValueError(f"No models with generateContent support found. Available models: {[m.name for m in genai.list_models()]}")
        
        # Initialize the model
        model = genai.GenerativeModel(model_name=model_name_to_use, generation_config=generation_config)
        
    except Exception as e:
        # If listing models fails, try common model names as fallback
        fallback_models = ["models/gemini-2.5-flash", "models/gemini-pro-latest", "models/gemini-flash-latest"]
        for fallback_name in fallback_models:
            try:
                model = genai.GenerativeModel(model_name=fallback_name, generation_config=generation_config)
                model_name_to_use = fallback_name
                break
            except:
                continue
        else:
            raise ValueError(f"Could not initialize any Gemini model. Tried to list models but got: {str(e)}. Fallback models also failed.")
    
    prompt = f"""You are an expert at extracting rubric information from documents. 

Analyze the following document and extract the grading rubric structure. The document may contain:
- A rubric table with criteria and point values
- A list of criteria with points
- Point ranges (like "0-1" or "0 – 0.5")
- Decimal point values (like 0.5, 1.0)

Extract the rubric information and return it as JSON in this exact format:
{{
  "name": "Rubric Name",
  "total_points": <total number or sum of all criteria>,
  "criteria": [
    {{
      "name": "Criterion Name",
      "max_points": <number>,
      "description": "Description or evaluation levels"
    }}
  ]
}}

IMPORTANT:
- If there are point ranges (e.g., "0 – 1"), use the maximum value as max_points
- If total points is mentioned (e.g., "out of 5 points"), use that; otherwise sum all criteria
- Include all criteria found in the rubric
- For descriptions, include the evaluation levels/descriptors if available
- Return ONLY valid JSON, no additional text or markdown

DOCUMENT:
{document_text}

Return the JSON now:"""
    
    try:
        response = model.generate_content(prompt)
        response_text = response.text.strip()
        
        # Extract JSON from markdown code blocks if present
        if "```json" in response_text:
            response_text = response_text.split("```json")[1].split("```")[0].strip()
        elif "```" in response_text:
            response_text = response_text.split("```")[1].split("```")[0].strip()
        
        # Find the first { and last } to extract just the JSON object
        first_brace = response_text.find("{")
        last_brace = response_text.rfind("}")
        
        if first_brace >= 0 and last_brace > first_brace:
            response_text = response_text[first_brace:last_brace + 1]
        elif first_brace >= 0:
            # Only opening brace found, try to balance braces
            brace_count = 0
            end_idx = first_brace
            for i in range(first_brace, len(response_text)):
                if response_text[i] == '{':
                    brace_count += 1
                elif response_text[i] == '}':
                    brace_count -= 1
                    if brace_count == 0:
                        end_idx = i + 1
                        break
            if end_idx > first_brace:
                response_text = response_text[first_brace:end_idx]
            else:
                raise ValueError(f"Incomplete JSON object (missing closing brace). Response (first 1000 chars): {response_text[:1000]}")
        else:
            raise ValueError(f"No JSON object found in response. Response (first 1000 chars): {response_text[:1000]}")
        
        # Parse JSON
        try:
            rubric_data = json.loads(response_text)
        except json.JSONDecodeError as json_err:
            # Provide detailed error information
            error_pos = getattr(json_err, 'pos', None)
            error_msg = getattr(json_err, 'msg', 'Unknown JSON error')
            error_details = f"JSON parse error"
            if error_pos:
                # Show context around the error
                start_context = max(0, error_pos - 50)
                end_context = min(len(response_text), error_pos + 50)
                context = response_text[start_context:end_context]
                error_details += f" at position {error_pos}: {error_msg}\nContext: ...{context}..."
            else:
                error_details += f": {error_msg}"
            raise ValueError(f"AI returned invalid JSON. {error_details}\nFull response length: {len(response_text)} chars\nFirst 1000 chars: {response_text[:1000]}")
        
        # Validate structure
        if 'name' not in rubric_data:
            rubric_data['name'] = 'Uploaded Rubric'
        if 'criteria' not in rubric_data or not isinstance(rubric_data['criteria'], list):
            raise ValueError("No criteria found in rubric")
        if len(rubric_data['criteria']) == 0:
            raise ValueError("Rubric must have at least one criterion")
        
        # Validate each criterion
        for criterion in rubric_data['criteria']:
            if 'name' not in criterion or 'max_points' not in criterion:
                raise ValueError("Each criterion must have name and max_points")
            if 'description' not in criterion:
                criterion['description'] = f"Evaluation of {criterion['name']}"
        
        # Calculate total points if not provided
        if rubric_data.get('total_points', 0) == 0:
            rubric_data['total_points'] = sum(c['max_points'] for c in rubric_data['criteria'])
        
        # Convert max_points to int if whole numbers
        for criterion in rubric_data['criteria']:
            max_pts = criterion['max_points']
            if isinstance(max_pts, (int, float)):
                if max_pts == int(max_pts):
                    criterion['max_points'] = int(max_pts)
                else:
                    criterion['max_points'] = float(max_pts)
        
        return rubric_data
        
    except ValueError as e:
        # Re-raise ValueError (which includes our detailed JSON errors)
        raise
    except Exception as e:
        raise ValueError(f"Error parsing rubric with AI: {str(e)}")


def parse_word_rubric_from_stream(file_stream):
    """
    Parse a Word document from a file stream to extract rubric information.
    Uses AI parsing exclusively for maximum flexibility.
    """
    if not DOCX_AVAILABLE:
        raise ValueError("python-docx library not installed. Install with: pip install python-docx")
    
    # Open document from stream (in memory)
    doc = Document(file_stream)
    
    # Extract all text from document
    text_content = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_content.append(para.text.strip())
    
    # Also extract text from tables
    if hasattr(doc, 'tables') and doc.tables:
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    table_text.append(row_text)
            if table_text:
                text_content.append('\n'.join(table_text))
    
    full_text = '\n'.join(text_content)
    
    if not full_text.strip():
        raise ValueError("Document appears to be empty")
    
    # Use AI parsing exclusively
    rubric_data = parse_rubric_with_ai(full_text)
    
    # Explicitly close/dereference the document to release file handle
    del doc
    
    return rubric_data


@grading_bp.route('/rubrics/upload', methods=['POST'])
def upload_rubric():
    """Upload a new rubric from JSON or Word document"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    filename_lower = file.filename.lower()
    
    # Check file type
    is_json = filename_lower.endswith('.json')
    is_word = filename_lower.endswith(('.doc', '.docx'))
    is_pdf = filename_lower.endswith('.pdf')
    
    if not (is_json or is_word or is_pdf):
        return jsonify({'error': 'File must be a JSON file (.json), Word document (.doc, .docx), or PDF (.pdf)'}), 400
    
    try:
        rubric_data = None
        original_filename = None
        original_file_content = None
        
        if is_json:
            # Handle JSON file
            content = file.read().decode('utf-8')
            rubric_data = json.loads(content)
        else:
            # Handle Word document or PDF - save original file
            from io import BytesIO
            file.seek(0)  # Reset file pointer
            original_file_content = file.read()
            original_filename = file.filename
            file_stream = BytesIO(original_file_content)
            
            if is_word:
                # Handle Word document
                if not DOCX_AVAILABLE:
                    return jsonify({'error': 'Word document support requires python-docx. Install with: pip install python-docx'}), 500
                
                # Parse directly from memory stream
                rubric_data = parse_word_rubric_from_stream(file_stream)
            elif is_pdf:
                # For PDF, we can't parse it automatically, but we can save it for embedding
                # User must provide rubric data separately or we'll create a placeholder
                # For now, return error suggesting to convert to Word or use JSON
                # But we could also allow PDF uploads for embedding-only use cases
                return jsonify({'error': 'PDF parsing not yet supported. Please upload a Word document (.docx) or JSON file. PDFs can be embedded but not automatically parsed.'}), 400
        
        # Validate rubric structure
        required_fields = ['name', 'total_points', 'criteria']
        for field in required_fields:
            if field not in rubric_data:
                return jsonify({'error': f'Missing required field: {field}'}), 400
        
        if not isinstance(rubric_data['criteria'], list) or len(rubric_data['criteria']) == 0:
            return jsonify({'error': 'Rubric must have at least one criterion'}), 400
        
        # Validate criteria structure
        for criterion in rubric_data['criteria']:
            if 'name' not in criterion or 'max_points' not in criterion:
                return jsonify({'error': 'Each criterion must have name and max_points'}), 400
            if 'description' not in criterion:
                criterion['description'] = f"Evaluation of {criterion['name']}"
        
        # Generate safe filename from rubric name
        rubric_name = rubric_data['name']
        safe_filename = "".join(c for c in rubric_name if c.isalnum() or c in (' ', '-', '_')).strip()
        safe_filename = safe_filename.replace(' ', '_').lower()
        if not safe_filename:
            safe_filename = 'uploaded_rubric'
        
        # Add timestamp to avoid conflicts
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        base_filename = f"{safe_filename}_{timestamp}"
        json_filename = f"{base_filename}.json"
        
        # Save to rubrics directory
        rubrics_dir = Path(__file__).parent.parent / 'rubrics'
        rubrics_dir.mkdir(exist_ok=True)
        json_path = rubrics_dir / json_filename
        
        # Store original filename in rubric metadata
        rubric_data['_original_filename'] = original_filename
        if original_filename:
            # Get original file extension
            original_ext = Path(original_filename).suffix.lower()
            rubric_data['_original_file_extension'] = original_ext
        
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(rubric_data, f, indent=2, ensure_ascii=False)
        
        # Save original file if it's not JSON
        if original_file_content and original_filename:
            original_ext = Path(original_filename).suffix.lower()
            original_filename_saved = f"{base_filename}{original_ext}"
            original_path = rubrics_dir / original_filename_saved
            with open(original_path, 'wb') as f:
                f.write(original_file_content)
            rubric_data['_original_file_path'] = original_filename_saved
        
        return jsonify({
            'success': True,
            'filename': json_filename,
            'name': rubric_name,
            'total_points': rubric_data['total_points'],
            'criteria_count': len(rubric_data['criteria']),
            'original_filename': original_filename_saved if original_file_content and original_filename else None,
            'message': f'Rubric "{rubric_name}" uploaded successfully'
        })
    
    except json.JSONDecodeError:
        return jsonify({'error': 'Invalid JSON format'}), 400
    except ValueError as e:
        error_msg = str(e)
        # Truncate very long error messages
        if len(error_msg) > 1000:
            error_msg = error_msg[:1000] + "..."
        return jsonify({'error': error_msg}), 400
    except Exception as e:
        return jsonify({'error': f'Error uploading rubric: {str(e)}'}), 500

